"""Convert extractable pilot source documents into validated local rules."""

from __future__ import annotations

import hashlib
import argparse
import json
import re
from datetime import datetime, timezone
from html.parser import HTMLParser
from pathlib import Path
from typing import Any

import fitz

from geode.schemas import ConfidenceScores, LayerIndexRecord, LocalRule, RuleUnit
from geode.utils.file_io import atomic_write_jsonl, iter_jsonl, load_json

LOCAL_LAYERS = {
    "county": "08_County_Authorities",
    "district": "09_District_Authorities",
}
CRS_PATTERN = re.compile(
    r"(?:C\.R\.S\.|CRS)[^0-9]{0,20}"
    r"(\d{1,2})-(\d+)-(\d+(?:\.\d+)?)",
    re.IGNORECASE,
)
EXCLUDED_DOCUMENT_NAMES = {"zone_lot_brochure.pdf"}
EXCLUDED_DOCUMENT_MARKERS = (
    "organizationalchart",
    "permitapplication",
    "candidateinformation",
    "holidaySchedule".casefold(),
    "boardsandcommissions",
    "brochure",
    "application",
    "booklet",
)


def ingest_local_rules(root: Path) -> dict[str, int]:
    """Extract, validate, index, and cross-reference pilot local rules."""

    resolved_root = root.resolve()
    manifest_rows = _download_rows(resolved_root)
    known_state_ids = _known_state_ids(resolved_root)
    authorities = _authority_lookup(resolved_root)
    source_registry = _source_registry_lookup(resolved_root)
    seen_hashes: set[str] = set()
    records_by_layer: dict[str, list[LocalRule]] = {layer: [] for layer in LOCAL_LAYERS.values()}
    units_by_layer: dict[str, list[RuleUnit]] = {layer: [] for layer in LOCAL_LAYERS.values()}
    quarantine: list[dict[str, Any]] = []

    for row in manifest_rows:
        source_entry = source_registry.get(str(row.get("source_id")), {})
        if str(row.get("authority_id")) not in authorities and source_entry:
            row = {
                **row,
                "authority_id": source_entry.get("authority_id"),
                "authority_level": source_entry.get("authority_level", row.get("authority_level")),
                "category": source_entry.get("category", row.get("category")),
            }
        raw_path = Path(str(row["raw_path"]))
        if raw_path.suffix.casefold() not in {".pdf", ".html", ".htm", ".txt", ".docx", ".bin"} or not raw_path.exists():
            continue
        if not row.get("category"):
            row = {**row, "category": _infer_category(raw_path)}
        if _is_excluded_document(raw_path):
            quarantine.append(
                _quarantine_row(
                    raw_path,
                    "Excluded non-rule administrative material; preserved for audit.",
                    row,
                )
            )
            continue
        source_hash = str(row.get("sha256") or _sha256(raw_path))
        if source_hash in seen_hashes:
            continue
        seen_hashes.add(source_hash)
        try:
            text, page_count, source_format = _extract_source(raw_path)
        except (AssertionError, OSError, ValueError, RuntimeError) as exc:
            quarantine.append(_quarantine_row(raw_path, f"Source extraction failed: {exc}", row))
            continue
        extracted_text = re.sub(r"\[Page \d+\]", "", text).strip()
        if len(extracted_text) < 500:
            quarantine.append(
                _quarantine_row(
                    raw_path,
                    "PDF text extraction returned insufficient text; OCR review is required.",
                    row,
                )
            )
            continue
        authority = authorities.get(str(row["authority_id"]))
        if authority is None:
            quarantine.append(
                _quarantine_row(
                    raw_path,
                    "Authority identity is not present in the local registry.",
                    row,
                )
            )
            continue
        rule = _build_rule(
            row,
            raw_path,
            text,
            page_count,
            source_hash,
            authority,
            known_state_ids,
            source_format,
        )
        layer = LOCAL_LAYERS[str(row["authority_level"])]
        units = _build_rule_units(rule, text)
        rule.rule_unit_ids = [unit.id for unit in units]
        records_by_layer[layer].append(rule)
        units_by_layer[layer].extend(units)

    for layer, records in records_by_layer.items():
        _write_layer_records(resolved_root, layer, records, units_by_layer[layer])
    quarantine_path = resolved_root / "_QUARANTINE" / "local_extraction_quarantine.jsonl"
    atomic_write_jsonl(quarantine_path, quarantine, resolved_root)
    return {
        "county_rules": len(records_by_layer[LOCAL_LAYERS["county"]]),
        "district_rules": len(records_by_layer[LOCAL_LAYERS["district"]]),
        "quarantined": len(quarantine),
    }


def _download_rows(root: Path) -> list[dict[str, Any]]:
    """Return successful local download rows whose files still exist."""

    path = root / "_CONTROL_PLANE" / "LOCAL_DOWNLOAD_MANIFEST.jsonl"
    rows: list[dict[str, Any]] = []
    for row in iter_jsonl(path):
        if row.get("status") != "downloaded":
            continue
        if Path(str(row.get("raw_path", ""))).exists():
            rows.append(row)
    return rows


def _is_excluded_document(path: Path) -> bool:
    """Exclude forms, brochures, and administrative material from legal rules."""

    name = path.name.casefold()
    compact_name = re.sub(r"[^a-z0-9]", "", name)
    return name in EXCLUDED_DOCUMENT_NAMES or any(marker in compact_name for marker in EXCLUDED_DOCUMENT_MARKERS)


def _authority_lookup(root: Path) -> dict[str, dict[str, Any]]:
    """Load pilot authority metadata from the source registry."""

    registry = load_json(root / "_CONTROL_PLANE" / "LOCAL_SOURCE_REGISTRY.json")
    pilot = registry.get("pilot", {})
    entries = [*pilot.get("counties", []), *pilot.get("districts", [])]
    return {str(entry["authority_id"]): entry for entry in entries}


def _source_registry_lookup(root: Path) -> dict[str, dict[str, Any]]:
    """Load current source metadata for identity repair of historical manifests."""

    registry = load_json(root / "_CONTROL_PLANE" / "LOCAL_SOURCE_REGISTRY.json")
    return {
        str(entry["source_id"]): entry
        for entry in registry.get("pilot", {}).get("county_sources", [])
    }


def _known_state_ids(root: Path) -> set[str]:
    """Read known CRS IDs from the state statute indexes."""

    ids: set[str] = set()
    for path in (root / "01_Statutes_CRS").rglob("*.jsonl"):
        for row in iter_jsonl(path):
            value = row.get("id") or row.get("entity_id")
            if isinstance(value, str) and value.startswith("CRS-"):
                ids.add(value)
    return ids


def _extract_source(path: Path) -> tuple[str, int, str]:
    """Extract text from supported source formats without changing raw files."""

    suffix = path.suffix.casefold()
    signature = path.read_bytes()[:16]
    if signature.startswith(b"%PDF"):
        document = fitz.open(path)
        pages = [f"[Page {index}]\n{page.get_text()}" for index, page in enumerate(document, start=1)]
        return "\n\n".join(pages), len(document), "pdf"
    if signature.startswith(b"PK") and suffix == ".bin":
        from docx import Document

        document = Document(path)
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        return f"[Page 1]\n{text}", 1, "docx"
    if b"ftyp" in signature[:12]:
        raise ValueError("binary media source is not a legal text document")
    if signature.lstrip().startswith((b"<html", b"<!doctype", b"<!DOCTYPE")):
        parser = _TextHTMLParser()
        parser.feed(path.read_text(encoding="utf-8", errors="replace"))
        return f"[Page 1]\n{parser.text}", 1, "html"
    if suffix in {".html", ".htm"}:
        parser = _TextHTMLParser()
        parser.feed(path.read_text(encoding="utf-8", errors="replace"))
        return f"[Page 1]\n{parser.text}", 1, "html"
    if suffix == ".txt":
        return f"[Page 1]\n{path.read_text(encoding='utf-8', errors='replace')}", 1, "text"
    if suffix == ".docx":
        from docx import Document

        document = Document(path)
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        return f"[Page 1]\n{text}", 1, "docx"
    if suffix == ".bin":
        raw = path.read_bytes()
        decoded = raw.decode("utf-8", errors="replace")
        if "<html" in decoded.casefold() or "<!doctype" in decoded.casefold():
            parser = _TextHTMLParser()
            parser.feed(decoded)
            return f"[Page 1]\n{parser.text}", 1, "html"
        if decoded.strip():
            return f"[Page 1]\n{decoded}", 1, "text"
        raise ValueError("binary source has no readable text signature")
    raise ValueError(f"unsupported local source format: {path.suffix}")


class _TextHTMLParser(HTMLParser):
    """Collect visible text from an HTML source page."""

    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []
        self._hidden_depth = 0

    @property
    def text(self) -> str:
        """Return normalized visible text."""

        return re.sub(r"\s+", " ", " ".join(self.parts)).strip()

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag.casefold() in {"script", "style", "noscript", "svg"}:
            self._hidden_depth += 1

    def handle_endtag(self, tag: str) -> None:
        if tag.casefold() in {"script", "style", "noscript", "svg"} and self._hidden_depth:
            self._hidden_depth -= 1

    def handle_data(self, data: str) -> None:
        if not self._hidden_depth and data.strip():
            self.parts.append(data.strip())


def _extract_pdf(path: Path) -> tuple[str, int]:
    """Extract page-marked text from a PDF without changing the raw file."""

    document = fitz.open(path)
    pages = [f"[Page {index}]\n{page.get_text()}" for index, page in enumerate(document, start=1)]
    return "\n\n".join(pages), len(document)


def _build_rule(
    row: dict[str, Any],
    raw_path: Path,
    text: str,
    page_count: int,
    source_hash: str,
    authority: dict[str, Any],
    known_state_ids: set[str],
    source_format: str,
) -> LocalRule:
    """Build one source-backed local rule record."""

    title = _title(raw_path, text)
    citations = sorted(
        {
            f"CRS-{title_num}-{article_num}-{section_num}"
            for title_num, article_num, section_num in CRS_PATTERN.findall(text)
        }
        & known_state_ids
    )
    citation_pages = _citation_pages(text, citations)
    section_heading, source_page, source_page_end = _provenance_span(text)
    source_line_start, source_line_end = _provenance_line_span(text)
    source_url = str(row["requested_url"])
    rule_id = "LOCAL-RULE-" + _safe_id(f"{row['authority_id']}-{raw_path.stem}-{source_hash[:10]}")
    effective_date = _effective_date(text)
    return LocalRule(
        id=rule_id,
        authority_id=str(row["authority_id"]),
        authority_level=str(row["authority_level"]),
        authority_type=str(authority["authority_type"]),
        authority_name=str(authority["name"]),
        district_family=authority.get("district_family"),
        county_names=[str(value) for value in authority.get("county_names", [])],
        citation=_citation(authority, raw_path, text),
        title=title,
        section_heading=section_heading,
        full_text=text.strip(),
        summary=_summary(title, text),
        state_authority_ids=citations,
        source_citation_pages=citation_pages,
        source_category=str(row.get("category")) if row.get("category") else None,
        source_format=source_format,
        provenance_status="section" if section_heading else "document",
        semantic_status="source_preservation_only",
        effective_date=effective_date,
        status="unknown",
        geographic_scope=[str(value) for value in authority.get("county_names", [])],
        source_url=source_url,
        source_path=raw_path.resolve().relative_to(Path.cwd().resolve()).as_posix(),
        source_hash=source_hash,
        source_version=source_hash,
        source_section=section_heading,
        source_page=source_page,
        source_page_end=source_page_end,
        source_line_start=source_line_start,
        source_line_end=source_line_end,
        data_retrieved=_retrieved_at(row),
        confidence=ConfidenceScores(
            overall=0.82 if page_count < 100 else 0.78,
            route="flag_accept",
            fields={"full_text": 0.9, "state_authority_ids": 0.95 if citations else 0.7},
        ),
    )


def _build_rule_units(rule: LocalRule, text: str) -> list[RuleUnit]:
    """Create conservative source-preservation units for explicit sections.

    These units do not interpret the law. They preserve section text and make
    semantic extraction status visible to downstream AI workflows.
    """

    spans = _section_spans(text)
    if not spans:
        spans = [("Document-level source", 1, 1, text)]
    units: list[RuleUnit] = []
    for index, (heading, start_page, end_page, body) in enumerate(spans, start=1):
        unit_id = f"{rule.id}-UNIT-{index:04d}"
        units.append(
            RuleUnit(
                id=unit_id,
                parent_regulation_id=rule.id,
                source_section=heading,
                rule_type="standard",
                regulated_entity="Not separately specified in source section",
                action_required=body.strip() or heading,
                plain_english_summary="Source section preserved; semantic extraction remains pending.",
                subject_tags=["compliance"],
                semantic_status="source_preservation_only",
                confidence=ConfidenceScores(
                    overall=0.45,
                    route="source_preservation_only",
                    fields={"source_section": 0.9, "atomicity": 0.2},
                ),
            )
        )
    return units


def _write_layer_records(
    root: Path,
    layer: str,
    records: list[LocalRule],
    units: list[RuleUnit],
) -> None:
    """Write local rules and merge their index rows with authority identities."""

    metadata_path = _write_with_version_fallback(
        root / layer / "_meta" / "local_rules.jsonl", records, root
    )
    units_path = _write_with_version_fallback(
        root / layer / "_meta" / "local_rule_units.jsonl", units, root
    )
    rule_rows = [_index_row(record, metadata_path, root, layer) for record in records]
    unit_rows = [_unit_index_row(unit, records, units_path, root, layer) for unit in units]
    index_path = root / layer / "_index.jsonl"
    existing = [
        row for row in iter_jsonl(index_path)
        if row.get("entity_type") not in {"local_rule", "rule_unit"}
    ]
    atomic_write_jsonl(index_path, [*existing, *rule_rows, *unit_rows], root)


def _write_with_version_fallback(target: Path, records: list[Any], root: Path) -> Path:
    """Write a generated file, using a versioned sibling if a sync lock persists."""

    try:
        atomic_write_jsonl(target, records, root)
        return target
    except PermissionError:
        version = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
        versioned = target.with_name(f"{target.stem}_{version}{target.suffix}")
        atomic_write_jsonl(versioned, records, root)
        return versioned


def _index_row(record: LocalRule, metadata_path: Path, root: Path, layer: str) -> LayerIndexRecord:
    """Create an AI discovery index row for one local rule."""

    relative_meta = metadata_path.resolve().relative_to(root.resolve()).as_posix()
    return LayerIndexRecord(
        id=record.id,
        layer=layer,
        entity_type=record.entity_type,
        title=record.title,
        citation=record.citation,
        path=relative_meta,
        meta_path=relative_meta,
        source_url=record.source_url,
        source_path=record.source_path,
        last_updated=record.data_retrieved,
        sha256=record.source_hash,
        tags=[record.authority_type, *record.county_names, *([record.source_category] if record.source_category else [])],
        confidence=record.confidence.overall,
        authority_id=record.authority_id,
        authority_name=record.authority_name,
        authority_level=record.authority_level,
        authority_type=record.authority_type,
        district_family=record.district_family,
        county_names=record.county_names,
        geographic_scope=record.geographic_scope,
        source_section=record.source_section,
        section_heading=record.section_heading,
        source_page=record.source_page,
        source_page_end=record.source_page_end,
        source_line_start=record.source_line_start,
        source_line_end=record.source_line_end,
        source_category=record.source_category,
        semantic_status=record.semantic_status,
        text_hash=hashlib.sha256(record.full_text.encode("utf-8")).hexdigest(),
    )


def _unit_index_row(
    unit: RuleUnit,
    records: list[LocalRule],
    units_path: Path,
    root: Path,
    layer: str,
) -> LayerIndexRecord:
    """Create an index row for a source-preservation unit."""

    parent_id = unit.parent_regulation_id
    parent = next(record for record in records if record.id == parent_id)
    relative_path = units_path.resolve().relative_to(root.resolve()).as_posix()
    return LayerIndexRecord(
        id=unit.id,
        layer=layer,
        entity_type=unit.entity_type,
        title=unit.source_section,
        citation=parent.citation,
        path=relative_path,
        meta_path=relative_path,
        source_url=parent.source_url,
        source_path=parent.source_path,
        last_updated=parent.data_retrieved,
        sha256=parent.source_hash,
        tags=[
            *parent.county_names,
            *([parent.source_category] if parent.source_category else []),
            "source_preservation_only",
        ],
        confidence=unit.confidence.overall,
        authority_id=parent.authority_id,
        authority_name=parent.authority_name,
        authority_level=parent.authority_level,
        authority_type=parent.authority_type,
        district_family=parent.district_family,
        county_names=parent.county_names,
        geographic_scope=parent.geographic_scope,
        source_section=unit.source_section,
        section_heading=unit.source_section,
        source_page=parent.source_page,
        source_page_end=parent.source_page_end,
        source_line_start=parent.source_line_start,
        source_line_end=parent.source_line_end,
        source_category=parent.source_category,
        semantic_status=unit.semantic_status,
        text_hash=hashlib.sha256(unit.action_required.encode("utf-8")).hexdigest(),
    )


def _title(path: Path, text: str) -> str:
    """Choose a source-derived title."""

    lines = [line.strip() for line in text.splitlines() if line.strip() and not line.startswith("[Page")]
    return " ".join(lines[:2])[:240] or path.stem.replace("_", " ")


def _summary(title: str, text: str) -> str:
    """Create a conservative source-derived summary."""

    body = " ".join(line.strip() for line in text.splitlines() if line.strip())
    return f"Source document titled {title}. Extracted text begins: {body[:420]}"


def _section_heading(text: str) -> str | None:
    """Return the first recognizable source heading."""

    for line in text.splitlines():
        clean = line.strip()
        if re.match(r"(?:ARTICLE|CHAPTER|PART|[IVX]+\.)", clean, re.IGNORECASE):
            return clean[:240]
    return None


def _section_spans(text: str) -> list[tuple[str, int, int, str]]:
    """Split page-marked text at explicit legal headings."""

    lines = text.splitlines()
    headings: list[tuple[int, str, int]] = []
    page = 1
    for index, line in enumerate(lines):
        page_match = re.fullmatch(r"\[Page (\d+)\]", line.strip())
        if page_match:
            page = int(page_match.group(1))
            continue
        clean = line.strip()
        if re.match(
            r"(?:CHAPTER|ARTICLE|PART|SECTION|Sec\.?|§|\d{1,4}[.-]\d+)",
            clean,
            re.IGNORECASE,
        ) and len(clean) >= 4:
            headings.append((index, clean[:240], page))
    spans: list[tuple[str, int, int, str]] = []
    for position, (line_index, heading, start_page) in enumerate(headings):
        end_index = headings[position + 1][0] if position + 1 < len(headings) else len(lines)
        end_page = start_page
        for line in lines[line_index:end_index]:
            match = re.fullmatch(r"\[Page (\d+)\]", line.strip())
            if match:
                end_page = int(match.group(1))
        body = "\n".join(lines[line_index:end_index]).strip()
        spans.append((heading, start_page, end_page, body))
    return spans


def _provenance_span(text: str) -> tuple[str | None, int, int]:
    """Return the first explicit section and its page range."""

    spans = _section_spans(text)
    if not spans:
        return None, 1, 1
    heading, start_page, end_page, _ = spans[0]
    return heading, start_page, end_page


def _provenance_line_span(text: str) -> tuple[int, int]:
    """Return the exact line range for the first section or full document."""

    lines = text.splitlines()
    headings = [
        index
        for index, line in enumerate(lines)
        if re.match(
            r"(?:CHAPTER|ARTICLE|PART|SECTION|Sec\\.?|Â§|\\d{1,4}[.-]\\d+)",
            line.strip(),
            re.IGNORECASE,
        )
    ]
    if not headings:
        headings = [
            index
            for index, line in enumerate(lines)
            if re.match(r"^(?:CHAPTER|ARTICLE|PART|SECTION|Sec|§|\d+[.-]\d+)", line.strip(), re.IGNORECASE)
        ]
    start = headings[0] if headings else 0
    end = headings[1] - 1 if len(headings) > 1 else max(len(lines) - 1, start)
    return start + 1, end + 1


def _citation_pages(text: str, citations: list[str]) -> dict[str, list[int]]:
    """Find the page numbers containing each normalized CRS citation."""

    pages: dict[str, list[int]] = {citation: [] for citation in citations}
    chunks = re.split(r"\[Page (\d+)\]", text)
    for index in range(1, len(chunks), 2):
        page_number = int(chunks[index])
        page_text = chunks[index + 1]
        for citation in citations:
            title, article, section = citation.removeprefix("CRS-").split("-")
            pattern = rf"(?:C\.R\.S\.|CRS)[^0-9]{{0,20}}{re.escape(title)}-{re.escape(article)}-{re.escape(section)}"
            if re.search(pattern, page_text, re.IGNORECASE):
                pages[citation].append(page_number)
    return {citation: values for citation, values in pages.items() if values}


def _citation(authority: dict[str, Any], path: Path, text: str) -> str:
    """Build a stable local citation from the authority and source document."""

    ordinance = re.search(r"ORDINANCE\s+NO\.\s+([0-9-]+)", text, re.IGNORECASE)
    if ordinance:
        return f"{authority['name']} Ordinance {ordinance.group(1)}"
    return f"{authority['name']} - {path.stem.replace('_', ' ')}"


def _effective_date(text: str):
    """Extract a clearly labeled effective date when present."""

    match = re.search(r"Effective(?: Date)?\s*[: ]+([A-Z][a-z]+\s+\d{1,2},\s+\d{4})", text)
    if not match:
        return None
    from datetime import datetime as DateTime

    try:
        return DateTime.strptime(match.group(1), "%B %d, %Y").date()
    except ValueError:
        return None


def _retrieved_at(row: dict[str, Any]) -> datetime:
    """Use the download timestamp when available."""

    value = str(row.get("retrieved_at") or "")
    return datetime.fromisoformat(value.replace("Z", "+00:00")) if value else datetime.now(timezone.utc)


def _quarantine_row(path: Path, reason: str, source_row: dict[str, Any]) -> dict[str, Any]:
    """Create a review record for a source that cannot be safely normalized."""

    source_hash = str(source_row.get("sha256") or _sha256(path))
    return {
        "event_id": f"QR-{datetime.now(timezone.utc).strftime('%Y%m%dT%H%M%SZ')}-LOCAL",
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "source_path": path.resolve().as_posix(),
        "source_id": source_row.get("source_id"),
        "requested_url": source_row.get("requested_url"),
        "authority_id": source_row.get("authority_id"),
        "authority_level": source_row.get("authority_level"),
        "source_category": source_row.get("category") or _infer_category(path),
        "source_hash": source_hash,
        "layer": "08_County_Authorities" if "county" in path.parts else "09_District_Authorities",
        "reason": reason,
        "confidence": 0.0,
        "reviewed": False,
    }


def _infer_category(path: Path) -> str:
    """Assign a transparent fallback category when historical rows lack one."""

    compact = re.sub(r"[^a-z0-9]", "", path.parent.name.casefold())
    known = {
        "ordinance": "county_ordinances",
        "code": "county_codes",
        "zoning": "land_use_zoning",
        "landuse": "land_use_zoning",
        "subdivision": "subdivision_development",
        "building": "building_construction",
        "publichealth": "public_health",
        "openburning": "environmental_open_burning",
        "transportation": "roads_transportation_access",
        "animalcontrol": "animal_control_nuisance",
        "emergency": "emergency_fire_restrictions",
        "resolution": "continuing_resolutions",
        "manual": "administrative_rule_manuals",
        "archive": "archived_versions",
    }
    for token, category in known.items():
        if token in compact:
            return category
    return "unclassified_local_source"


def _sha256(path: Path) -> str:
    """Hash a raw source file."""

    return hashlib.sha256(path.read_bytes()).hexdigest()


def _safe_id(value: str) -> str:
    """Create a stable uppercase identifier fragment."""

    return re.sub(r"[^A-Za-z0-9]+", "_", value).strip("_").upper()


def main() -> int:
    """Run local-source normalization from the command line."""

    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--root", type=Path, default=Path.cwd())
    args = parser.parse_args()
    print(json.dumps(ingest_local_rules(args.root.resolve()), indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
